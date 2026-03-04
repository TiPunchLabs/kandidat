"""CV format conversion service: HTML to PDF (WeasyPrint) and HTML to DOCX (python-docx).

PDF: Injects print-safe CSS (A4, margins, color preservation) for faithful rendering.
DOCX: Uses BeautifulSoup semantic parsing for structured, ATS-optimized output.

Adapted from the cv-converter project (https://github.com/TiPunchLabs/cv-converter).
"""

import logging
import re
from io import BytesIO
from pathlib import Path

logger = logging.getLogger(__name__)

# CSS injected into every PDF render for faithful A4 output
_PRINT_CSS = """
@page {
    size: A4;
    margin: 0;
}
body {
    margin: 0;
    padding: 0;
    zoom: 1;
    -webkit-print-color-adjust: exact;
    print-color-adjust: exact;
}
.print-button {
    display: none !important;
}
"""

# Section title translations (FR/EN)
_TRANSLATIONS = {
    "fr": {
        "profil": "PROFIL PROFESSIONNEL",
        "skills": "COMPÉTENCES TECHNIQUES",
        "certs": "CERTIFICATIONS PROFESSIONNELLES",
        "experience": "EXPÉRIENCE PROFESSIONNELLE",
        "education": "FORMATION",
        "projects": "PROJETS GITHUB",
    },
    "en": {
        "profil": "PROFESSIONAL PROFILE",
        "skills": "TECHNICAL SKILLS",
        "certs": "PROFESSIONAL CERTIFICATIONS",
        "experience": "PROFESSIONAL EXPERIENCE",
        "education": "EDUCATION",
        "projects": "GITHUB PROJECTS",
    },
}


def validate_html(html_content: str) -> bool:
    """Check that HTML content is non-empty and minimally well-formed."""
    if not html_content or not html_content.strip():
        return False
    return "<" in html_content and ">" in html_content


def detect_language(html_content: str) -> str:
    """Detect language from HTML content. Returns 'fr' or 'en'."""
    # Check <html lang="xx"> attribute
    match = re.search(r'<html[^>]+lang=["\'](\w+)["\']', html_content)
    if match:
        lang = match.group(1).lower()
        if lang.startswith("en"):
            return "en"
        if lang.startswith("fr"):
            return "fr"

    # Check for French keywords
    if any(word in html_content for word in ["Compétences", "Expérience", "Formation"]):
        return "fr"
    if any(word in html_content for word in ["Skills", "Experience", "Education"]):
        return "en"

    return "fr"


def html_to_pdf(html_content: str, base_url: str | None = None) -> bytes:
    """Convert HTML string to PDF bytes using WeasyPrint with print-safe CSS.

    Args:
        html_content: The HTML string to convert.
        base_url: Optional base URL for resolving relative resources (fonts, images).
    """
    from weasyprint import CSS, HTML

    if not validate_html(html_content):
        raise ValueError("Contenu HTML invalide pour la conversion PDF")

    additional_css = CSS(string=_PRINT_CSS)
    html = HTML(string=html_content, base_url=base_url)
    return html.write_pdf(stylesheets=[additional_css])


def html_to_docx(html_content: str) -> bytes:
    """Convert HTML string to DOCX bytes using BeautifulSoup semantic parsing.

    Extracts structured CV sections (header, profile, skills, experience, education,
    projects, certifications) and reconstructs them as a styled, ATS-optimized DOCX.
    """
    from bs4 import BeautifulSoup
    from docx import Document
    from docx.shared import Inches

    if not validate_html(html_content):
        raise ValueError("Contenu HTML invalide pour la conversion DOCX")

    soup = BeautifulSoup(html_content, "html.parser")
    language = detect_language(html_content)
    doc = Document()

    # Configure margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Try semantic extraction first; fall back to generic parsing
    if _has_semantic_structure(soup):
        _build_semantic_docx(doc, soup, language)
    else:
        _build_generic_docx(doc, soup)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def save_cv_files(html_content: str, slug: str, prefix: str = "cv_adapte") -> list[dict]:
    """Generate PDF and DOCX from HTML and save them as candidature files.

    Returns list of dicts with file info: [{"nom": ..., "type": ...}, ...]
    Handles naming collisions by adding numeric suffix.
    """
    from services.database import Candidature, Fichier, db

    candidature = db.session.get(Candidature, slug)
    if candidature is None:
        raise ValueError(f"Candidature '{slug}' non trouvee")

    data_dir = _data_dir()
    cand_dir = data_dir / "candidatures" / slug
    cand_dir.mkdir(parents=True, exist_ok=True)

    results = []

    for ext, converter in [("pdf", html_to_pdf), ("docx", html_to_docx)]:
        content = converter(html_content)
        base_name = f"{prefix}_{slug}"
        final_name = _unique_filename(slug, base_name, ext)

        dest = cand_dir / final_name
        dest.write_bytes(content)

        chemin = f"candidatures/{slug}/{final_name}"
        fichier = Fichier(slug=slug, nom=final_name, chemin=chemin, type=ext)
        db.session.add(fichier)
        results.append({"nom": final_name, "type": ext})

    db.session.commit()
    return results


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _unique_filename(slug: str, base_name: str, ext: str) -> str:
    """Generate a unique filename, adding numeric suffix if needed."""
    from services.database import Fichier

    candidate = f"{base_name}.{ext}"
    existing = Fichier.query.filter_by(slug=slug, nom=candidate).first()
    if not existing:
        return candidate

    counter = 2
    while True:
        candidate = f"{base_name}_{counter}.{ext}"
        existing = Fichier.query.filter_by(slug=slug, nom=candidate).first()
        if not existing:
            return candidate
        counter += 1


def _data_dir() -> Path:
    from flask import current_app

    return Path(current_app.config["FT_DATA_DIR"])


def _has_semantic_structure(soup) -> bool:
    """Check if the HTML has CV-specific semantic classes."""
    return bool(
        soup.find("div", class_="header")
        or soup.find("div", class_="experience-item")
        or soup.find("div", class_="skills-grid")
        or soup.find("div", class_="skill-bar-container")
    )


# ---------------------------------------------------------------------------
# Semantic DOCX builder (structured CV with known CSS classes)
# ---------------------------------------------------------------------------


def _build_semantic_docx(doc, soup, language: str):
    """Build a styled DOCX from a semantically-structured CV HTML."""
    tr = _TRANSLATIONS.get(language, _TRANSLATIONS["fr"])

    _add_header(doc, soup)
    _add_profile(doc, soup, tr)
    _add_skills(doc, soup, tr)
    _add_certifications(doc, soup, tr)
    _add_experience(doc, soup, tr)
    _add_education(doc, soup, tr)
    _add_projects(doc, soup, tr)


def _add_header(doc, soup):
    """Add styled header: name, subtitle, contact info, separator."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    header = soup.find("div", class_="header")
    if not header:
        return

    h1 = header.find("h1")
    if h1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h1.text.strip())
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 64, 175)

    subtitle_div = header.find("div", class_="subtitle")
    if subtitle_div:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle_div.text.strip())
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(59, 130, 246)

    # Contact info: in header or document root
    contact_div = header.find("div", class_="contact-info")
    if not contact_div:
        contact_div = soup.find("div", class_="contact-info")
    if contact_div:
        contact = contact_div.get_text(separator=" | ", strip=True)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(contact)
        run.font.size = Pt(10)

    _add_horizontal_line(doc)
    doc.add_paragraph()


def _add_profile(doc, soup, tr: dict):
    """Add professional profile section."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    profile_text = soup.find("p", class_="profile-text")
    if not profile_text:
        return

    _add_section_title(doc, tr["profil"])
    p = doc.add_paragraph(profile_text.text.strip())
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(10)
    doc.add_paragraph()


def _add_skills(doc, soup, tr: dict):
    """Add technical skills section (supports classic grid and design bar formats)."""
    from docx.shared import Pt, RGBColor

    skills_added = False

    # Format 1: classic skills-grid
    skills_grid = soup.find("div", class_="skills-grid")
    if skills_grid:
        _add_section_title(doc, tr["skills"])
        for skill_item in skills_grid.find_all("div", class_="skill-item"):
            category = skill_item.find("span", class_="skill-category")
            skill_list = skill_item.find("span", class_="skill-list")
            if category and skill_list:
                p = doc.add_paragraph()
                run = p.add_run(category.text.strip() + ": ")
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(30, 64, 175)
                run = p.add_run(skill_list.text.strip())
                run.font.size = Pt(10)
        skills_added = True

    # Format 2: design skill-bar-container
    if not skills_added:
        skill_bars = soup.find_all("div", class_="skill-bar-container")
        if skill_bars:
            _add_section_title(doc, tr["skills"])
            for skill_bar in skill_bars:
                label = skill_bar.find("div", class_="skill-bar-label")
                if label:
                    spans = label.find_all("span")
                    if len(spans) >= 2:
                        p = doc.add_paragraph()
                        run = p.add_run(f"{spans[0].text.strip()}: ")
                        run.font.bold = True
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(30, 64, 175)
                        run = p.add_run(spans[1].text.strip())
                        run.font.size = Pt(10)
            skills_added = True

    if skills_added:
        doc.add_paragraph()


def _add_certifications(doc, soup, tr: dict):
    """Add certifications section."""
    from docx.shared import Pt

    cert_list = soup.find("ul", class_="cert-list")
    if not cert_list:
        return

    _add_section_title(doc, tr["certs"])
    for cert in cert_list.find_all("li"):
        text = cert.get_text(strip=True)
        if not text.startswith("\u2713"):
            text = "\u2713 " + text
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        run.font.size = Pt(10)
    doc.add_paragraph()


def _add_experience(doc, soup, tr: dict):
    """Add professional experience section."""
    from docx.shared import Inches, Pt, RGBColor

    experiences = soup.find_all("div", class_="experience-item")
    if not experiences:
        return

    _add_section_title(doc, tr["experience"])
    for exp in experiences:
        job_title = exp.find("div", class_="job-title")
        if job_title:
            p = doc.add_paragraph()
            run = p.add_run(job_title.text.strip())
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(30, 64, 175)

        company = exp.find("div", class_="company-name")
        if company:
            p = doc.add_paragraph(company.text.strip())
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.italic = True

        date_range = exp.find("div", class_="date-range")
        if date_range:
            p = doc.add_paragraph(date_range.text.strip())
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(156, 163, 175)

        responsibilities = exp.find("ul", class_="responsibilities")
        if responsibilities:
            for li in responsibilities.find_all("li"):
                p = doc.add_paragraph(li.text.strip(), style="List Bullet")
                for run in p.runs:
                    run.font.size = Pt(10)

        tech_stack = exp.find("div", class_="tech-stack")
        if tech_stack:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(tech_stack.get_text(strip=True))
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(59, 130, 246)

        doc.add_paragraph()


def _add_education(doc, soup, tr: dict):
    """Add education section."""
    from docx.shared import Pt

    education_items = soup.find_all("div", class_="education-item")
    if not education_items:
        return

    _add_section_title(doc, tr["education"])
    for edu in education_items:
        job_title = edu.find("div", class_="job-title")
        if job_title:
            p = doc.add_paragraph()
            run = p.add_run(job_title.text.strip())
            run.font.bold = True
            run.font.size = Pt(11)

        company = edu.find("div", class_="company-name")
        if company:
            p = doc.add_paragraph(company.text.strip())
            for run in p.runs:
                run.font.size = Pt(10)

        date_range = edu.find("div", class_="date-range")
        if date_range:
            p = doc.add_paragraph(date_range.text.strip())
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.italic = True

        responsibilities = edu.find("ul", class_="responsibilities")
        if responsibilities:
            for li in responsibilities.find_all("li"):
                p = doc.add_paragraph(li.text.strip(), style="List Bullet")
                for run in p.runs:
                    run.font.size = Pt(9)

        doc.add_paragraph()


def _add_projects(doc, soup, tr: dict):
    """Add projects section (repo-card, projects-list, or projects-grid)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    projects_added = False

    # Format 1: repo-card
    repo_cards = soup.find_all("div", class_="repo-card")
    if repo_cards:
        _add_section_title(doc, tr["projects"])
        for card in repo_cards:
            link = card.find("a")
            if link:
                p = doc.add_paragraph()
                run = p.add_run(link.get_text(strip=True))
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(30, 64, 175)

            desc = card.find("div", class_="repo-description")
            if desc:
                p = doc.add_paragraph()
                run = p.add_run(desc.get_text(strip=True))
                run.font.size = Pt(9)
                run.font.italic = True

            tech = card.find("div", class_="repo-tech")
            if tech:
                tech_text = ", ".join([span.get_text(strip=True) for span in tech.find_all("span")])
                if tech_text:
                    p = doc.add_paragraph()
                    run = p.add_run(tech_text)
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(59, 130, 246)
        projects_added = True

    # Format 2: projects-list
    if not projects_added:
        projects_list = soup.find("div", class_="projects-list")
        if projects_list:
            _add_section_title(doc, tr["projects"])
            text = projects_list.get_text(separator="\n", strip=True)
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.size = Pt(10)
            projects_added = True

    # Format 3: projects-grid
    if not projects_added:
        projects_grid = soup.find("div", class_="projects-grid")
        if projects_grid:
            _add_section_title(doc, tr["projects"])
            text = projects_grid.get_text(separator="\n", strip=True)
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.size = Pt(10)
            projects_added = True

    # GitHub profile link
    if projects_added:
        github_profile = soup.find("div", class_="github-profile")
        if github_profile:
            link = github_profile.find("a")
            if link:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"Portfolio: {link.get('href', link.get_text(strip=True))}")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(59, 130, 246)


def _add_section_title(doc, title: str):
    """Add a styled section title with underline."""
    from docx.shared import Pt, RGBColor

    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 64, 175)

    _add_horizontal_line(doc, "3b82f6")


def _add_horizontal_line(doc, color: str = "3b82f6"):
    """Add a horizontal line below the current paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)

    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)

    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Generic DOCX builder (fallback for non-semantic HTML)
# ---------------------------------------------------------------------------


def _build_generic_docx(doc, soup):
    """Fallback: extract text content by HTML tags into a simple DOCX."""
    from docx.shared import Pt

    heading_tags = {"h1", "h2", "h3", "h4", "h5", "h6"}

    for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "div"]):
        text = element.get_text(strip=True)
        if not text:
            continue

        # Skip nested divs that would duplicate content
        if element.name == "div" and element.find(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li"]):
            continue

        tag = element.name.lower()
        if tag in heading_tags:
            level = int(tag[1])
            heading = doc.add_heading(text, level=level)
            for run in heading.runs:
                run.font.size = Pt(max(11, 20 - (level * 2)))
        elif tag == "li":
            p = doc.add_paragraph(text, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(10)
        else:
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.size = Pt(10)
